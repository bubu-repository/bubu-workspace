#!/usr/bin/env python3
"""
BUBU Minutes of Meeting (MOM) generator.

Reads a structured JSON file describing a meeting and produces a branded
.docx (and, when LibreOffice is available, a matching .pdf).

Usage:
    python3 generate_mom.py mom_data.json [output_basename]

The JSON schema is documented in references/data_schema.md. Every field is
optional except `meeting_title`; missing sections are simply skipped so the
same engine works for a 3-line internal standup and a full client review.
"""

import json
import os
import re
import subprocess
import sys
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# BUBU brand system
# ---------------------------------------------------------------------------
ORANGE = RGBColor(0xFF, 0x59, 0x00)      # BUBU signature orange
INK = RGBColor(0x14, 0x14, 0x16)         # near-black for body text
MUTED = RGBColor(0x70, 0x70, 0x74)       # secondary / labels
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HAIRLINE = "E3E3E3"
LIGHT = "F5F4F2"                          # warm light fill for meta / zebra
ORANGE_HEX = "FF5900"
INK_HEX = "141416"

DISPLAY = "Bebas Neue"                    # headlines / section titles
BODY = "Arial"                            # body copy (Helvetica/Liberation fallback)

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.normpath(os.path.join(HERE, "..", "assets"))
LOGO = os.path.join(ASSETS, "bubu_logo.png")
FONT_DIR = os.path.join(ASSETS, "fonts")

TAGLINE = "BUBU · 30 years turning culture into brands, products & experiences"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def _shade(element, fill_hex):
    """Apply a solid background fill to a cell or paragraph."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def shade_cell(cell, fill_hex):
    _shade(cell._tc.get_or_add_tcPr(), fill_hex)


def shade_paragraph(par, fill_hex):
    _shade(par._p.get_or_add_pPr(), fill_hex)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                     ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def set_table_borders(table, color=HAIRLINE, size=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def cell_bottom_border(cell, color=ORANGE_HEX, size=18):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:color"), color)
    borders.append(b)
    tcPr.append(borders)


def style_run(run, font=BODY, size=10.5, color=INK, bold=False, caps=False, spacing=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.all_caps = caps
    # ensure east-asian / complex fonts also map (helps LibreOffice)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)
    if spacing is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rpr.append(sp)
    return run


def add_para(container, text="", font=BODY, size=10.5, color=INK, bold=False,
             caps=False, align=None, before=0, after=4, line=1.12, spacing=None):
    """Add a paragraph either to the document or inside a table cell."""
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line
    if align is not None:
        par.alignment = align
    if text != "":
        style_run(par.add_run(text), font, size, color, bold, caps, spacing)
    return par


# ---------------------------------------------------------------------------
# Building blocks
# ---------------------------------------------------------------------------
def first_cell_clear(cell):
    """Remove the default empty paragraph python-docx puts in a fresh cell."""
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)


def add_header_band(doc, data):
    """Logo row + orange title band + meeting type tag."""
    content_w = Cm(17.0)  # A4 width 21cm - 2x2cm margins

    # --- Row 1: logo (left) + meeting-type pill (right) ---
    top = doc.add_table(rows=1, cols=2)
    top.allow_autofit = False
    no_table_borders(top)
    top.columns[0].width = Cm(11.0)
    top.columns[1].width = Cm(6.0)

    logo_cell = top.cell(0, 0)
    first_cell_clear(logo_cell)
    set_cell_margins(logo_cell, 0, 0, 0, 0)
    lp = logo_cell.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO):
        lp.add_run().add_picture(LOGO, height=Cm(0.95))

    tag_cell = top.cell(0, 1)
    first_cell_clear(tag_cell)
    tag_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(tag_cell, 0, 0, 0, 0)
    mtype = (data.get("meeting_type") or "Internal").strip().upper()
    tp = tag_cell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tp.paragraph_format.space_after = Pt(0)
    style_run(tp.add_run(f"{mtype} MEETING"), DISPLAY, 13, ORANGE, spacing=20)

    # breathing room between the logo and the orange title band
    spacer = add_para(doc, "", before=0, after=0)
    spacer.paragraph_format.line_spacing = Pt(10)

    # --- Row 2: orange title band ---
    band = doc.add_table(rows=1, cols=1)
    band.allow_autofit = False
    no_table_borders(band)
    band.columns[0].width = content_w
    bc = band.cell(0, 0)
    first_cell_clear(bc)
    bc.width = content_w
    shade_cell(bc, ORANGE_HEX)
    set_cell_margins(bc, 150, 150, 220, 160)

    title = data.get("document_title", "Minutes of Meeting")
    tpar = bc.add_paragraph()
    tpar.paragraph_format.space_after = Pt(0)
    tpar.paragraph_format.line_spacing = 0.95
    style_run(tpar.add_run(title.upper()), DISPLAY, 30, WHITE, spacing=20)

    # (subtitle line under the title removed per BUBU template — the meeting
    # topic lives in the metadata block below, not in the orange band.)

    add_para(doc, "", after=6)


def add_meta_block(doc, data):
    """Two-column key/value grid for meeting logistics."""
    pairs = []

    def add(label, key):
        v = data.get(key)
        if v:
            pairs.append((label, str(v)))

    # Category leads the block. Drop a trailing "Meeting" from the value since the
    # block is already about a meeting (e.g. "Discovery & Strategic Meeting" shows
    # as "Discovery & Strategic"). Project is intentionally NOT shown here (it's
    # only used for auto-filing) to avoid repeating it next to the meeting topic.
    cat = (data.get("category") or "").strip()
    if cat:
        cat = re.sub(r"\s*Meeting\s*$", "", cat).strip() or cat
        pairs.append(("Category", cat))
    if data.get("meeting_title"):
        pairs.append(("Meeting", str(data["meeting_title"])))
    add("Date", "date")
    add("Time", "time")
    add("Location / Platform", "location")
    add("Client", "client")
    add("Reference No.", "reference")
    add("Facilitator", "facilitator")
    add("Minute Taker", "minute_taker")
    add("Prepared by", "prepared_by")

    if not pairs:
        return

    # render two label/value pairs per row
    rows = (len(pairs) + 1) // 2
    tbl = doc.add_table(rows=rows, cols=4)
    tbl.allow_autofit = False
    no_table_borders(tbl)
    widths = [Cm(3.4), Cm(5.1), Cm(3.4), Cm(5.1)]
    for r in range(rows):
        for c in range(4):
            tbl.cell(r, c).width = widths[c]

    for i, (label, value) in enumerate(pairs):
        r = i // 2
        base = (i % 2) * 2
        lc = tbl.cell(r, base)
        vc = tbl.cell(r, base + 1)
        first_cell_clear(lc)
        first_cell_clear(vc)
        set_cell_margins(lc, 70, 70, 40, 80)
        set_cell_margins(vc, 70, 70, 40, 120)
        shade_cell(lc, LIGHT)
        shade_cell(vc, LIGHT)
        lp = lc.add_paragraph()
        lp.paragraph_format.space_after = Pt(0)
        style_run(lp.add_run(label.upper()), BODY, 8, MUTED, bold=True, spacing=8)
        vp = vc.add_paragraph()
        vp.paragraph_format.space_after = Pt(0)
        style_run(vp.add_run(value), BODY, 10, INK)

    # tidy a trailing empty pair
    if len(pairs) % 2 == 1:
        for c in (2, 3):
            cell = tbl.cell(rows - 1, c)
            shade_cell(cell, LIGHT)
    # no trailing spacer here — the first section header's space_before owns the gap


# Uniform spacing around every section header (in points).
SECTION_GAP_BEFORE = 13   # space above the header (from previous section's content)
SECTION_GAP_AFTER = 6     # space below the header (to its own content)


def add_section_header(doc, label, keep_next=True):
    """Orange square + Bebas Neue uppercase title with an orange underline.

    Paragraph-based (not a table) so spacing above/below is precise and identical
    for every section. Content renderers zero their *trailing* space, so the gap
    above the next header is always exactly SECTION_GAP_BEFORE — uniform whether
    the previous section ended in a table, bullets, or a paragraph.
    """
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(SECTION_GAP_BEFORE)
    pf.space_after = Pt(SECTION_GAP_AFTER)
    pf.keep_with_next = keep_next  # never leave a header stranded at a page bottom
    cell_para_bottom_border(par, ORANGE_HEX, 16)
    style_run(par.add_run("■  "), DISPLAY, 14, ORANGE)
    style_run(par.add_run(label.upper()), DISPLAY, 15, INK, spacing=14)
    return par


def add_paragraph_text(doc, text):
    blocks = [b.strip() for b in str(text).split("\n") if b.strip()]
    for i, block in enumerate(blocks):
        # last paragraph carries no trailing space so the gap to the next
        # section header is governed solely by the header's space_before
        after = 0 if i == len(blocks) - 1 else 8
        add_para(doc, block, BODY, 10.5, INK, after=after, line=1.4)


def add_bullets(doc, items, numbered=False):
    n = len(items)
    for i, item in enumerate(items, 1):
        if isinstance(item, dict):
            head = item.get("title") or item.get("point") or ""
            body = item.get("detail") or item.get("note") or ""
        else:
            head, body = str(item), ""
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.7)
        # last bullet carries no trailing space (uniform gap to next header)
        par.paragraph_format.space_after = Pt(0 if i == n else 5)
        par.paragraph_format.line_spacing = 1.32
        marker = f"{i}." if numbered else "•"  # round bullet (no em dash)
        style_run(par.add_run(f"{marker}  "), BODY, 10.5, ORANGE, bold=True)
        if head and body:
            style_run(par.add_run(f"{head}. "), BODY, 10.5, INK, bold=True)
            style_run(par.add_run(body), BODY, 10.5, INK)
        else:
            style_run(par.add_run(head or body), BODY, 10.5, INK)


def cell_borders(cell, color=HAIRLINE, size=4, left_color=None, left_size=None):
    """Set all four cell borders, optionally with a distinct (thicker) left edge."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        if edge == "left" and left_color:
            e.set(qn("w:sz"), str(left_size or size))
            e.set(qn("w:color"), left_color)
        else:
            e.set(qn("w:sz"), str(size))
            e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def add_highlight_box(doc, text):
    """A shaded, bordered callout with italic text — used inside a Discussion point
    to spotlight a key agreement or conclusion. Use sparingly: only when a point
    genuinely has something worth pulling out of the bullets.
    """
    add_para(doc, "", after=2)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.allow_autofit = False
    no_table_borders(tbl)
    tbl.columns[0].width = Cm(17.0)
    cell = tbl.cell(0, 0)
    first_cell_clear(cell)
    cell.width = Cm(17.0)
    shade_cell(cell, LIGHT)
    set_cell_margins(cell, 130, 130, 200, 160)
    # hairline all round + a bold orange left accent bar (on brand)
    cell_borders(cell, color="DEDDD8", size=4, left_color=ORANGE_HEX, left_size=26)
    blocks = [b.strip() for b in str(text).split("\n") if b.strip()]
    for i, block in enumerate(blocks):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0 if i == len(blocks) - 1 else 4)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(block)
        style_run(run, BODY, 10.5, INK)
        run.font.italic = True
    add_para(doc, "", after=6)


def render_flow_png(steps, out_path):
    """Draw a simple, on-brand process-flow infographic (orange rounded boxes with
    arrows) for an 'alur kegiatan / progress' diagram. Wraps into a serpentine
    layout (left-to-right, then right-to-left) so it stays readable for any count.
    Returns the natural (width, height) in pixels at 96 dpi, or None on failure.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:  # noqa: BLE001
        return None

    steps = [str(s) for s in steps if str(s).strip()]
    if not steps:
        return None

    S = 3  # supersample for crisp text
    per_row = len(steps) if len(steps) <= 4 else 4
    rows = (len(steps) + per_row - 1) // per_row
    boxW, boxH = 286 * S, 92 * S
    gapX, gapY = 64 * S, 60 * S
    margin = 18 * S
    cols = min(per_row, len(steps))
    W = margin * 2 + cols * boxW + (cols - 1) * gapX
    H = margin * 2 + rows * boxH + (rows - 1) * gapY

    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)

    def load(sz):
        for p in ("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                  "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
                  "/Library/Fonts/Arial Bold.ttf"):
            if os.path.exists(p):
                return ImageFont.truetype(p, sz)
        return ImageFont.load_default()

    font = load(23 * S)
    orange = (255, 89, 0)
    ink = (37, 37, 43)
    white = (255, 255, 255)

    pos = []
    for i in range(len(steps)):
        r = i // per_row
        pir = i % per_row
        col = pir if r % 2 == 0 else (per_row - 1 - pir)
        pos.append((margin + col * (boxW + gapX), margin + r * (boxH + gapY)))

    def arrow(x1, y1, x2, y2):
        import math
        d.line((x1, y1, x2, y2), fill=ink, width=3 * S)
        ang = math.atan2(y2 - y1, x2 - x1)
        l = 11 * S
        for s in (-0.45, 0.45):
            d.line((x2, y2, x2 - l * math.cos(ang + s), y2 - l * math.sin(ang + s)),
                   fill=ink, width=3 * S)

    for i in range(len(steps) - 1):
        x1, y1 = pos[i]
        x2, y2 = pos[i + 1]
        if (i // per_row) == ((i + 1) // per_row):  # same row
            if x2 > x1:
                arrow(x1 + boxW, y1 + boxH // 2, x2 - 4 * S, y2 + boxH // 2)
            else:
                arrow(x1, y1 + boxH // 2, x2 + boxW + 4 * S, y2 + boxH // 2)
        else:  # wrap: vertical connector (same column in serpentine)
            arrow(x1 + boxW // 2, y1 + boxH, x2 + boxW // 2, y2 - 4 * S)

    for i, (x, y) in enumerate(pos):
        d.rounded_rectangle((x, y, x + boxW, y + boxH), radius=18 * S, fill=orange)
        words = steps[i].split()
        lines, cur = [], ""
        for w in words:
            t = (cur + " " + w).strip()
            if d.textlength(t, font=font) <= boxW - 26 * S:
                cur = t
            else:
                if cur:
                    lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
        bb = font.getbbox("Ay")
        lh = (bb[3] - bb[1]) + 7 * S
        ty = y + (boxH - lh * len(lines)) // 2
        for ln in lines:
            tw = d.textlength(ln, font=font)
            d.text((x + (boxW - tw) // 2, ty), ln, fill=white, font=font)
            ty += lh

    img.save(out_path)
    return W // S, H // S


def add_flow(doc, steps):
    """Render a flow diagram and embed it; fall back to an arrow text line if PIL
    is unavailable so content is never lost."""
    import tempfile
    fd, path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    try:
        dims = render_flow_png(steps, path)
        if not dims:
            add_para(doc, "  →  ".join(str(s) for s in steps), BODY, 10.5, INK, after=8)
            return
        w, _ = dims
        width_cm = min(16.5, max(8.0, w / 96 * 2.54))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(path, width=Cm(width_cm))
    finally:
        try:
            os.remove(path)
        except OSError:
            pass


def add_discussion_image(doc, image_path):
    """Embed a user-supplied infographic/diagram image in a discussion point."""
    if not image_path or not os.path.exists(image_path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    try:
        p.add_run().add_picture(image_path, width=Cm(16.5))
    except Exception:  # noqa: BLE001
        p.add_run().add_picture(image_path)


def add_discussion(doc, items):
    for idx, item in enumerate(items):
        if isinstance(item, dict):
            topic = item.get("topic") or item.get("title") or ""
            intro = item.get("intro") or ""
            lead = item.get("lead") or ""
            notes = item.get("notes") or item.get("detail") or item.get("summary") or ""
            flow = item.get("flow") or ""
            image = item.get("image") or item.get("diagram_image") or ""
            highlight = item.get("highlight") or item.get("callout") or ""
        else:
            topic, intro, lead, notes, flow, image, highlight = "", "", "", str(item), "", "", ""
        if topic:
            # first topic sits right under the section header (header already
            # added its gap); later topics get a clear gap to separate blocks
            before = 2 if idx == 0 else 14
            tp = add_para(doc, "", before=before, after=5)
            style_run(tp.add_run(topic), BODY, 11.5, ORANGE, bold=True)
        if intro:
            add_paragraph_text(doc, intro)
        if lead:
            add_para(doc, lead, BODY, 10.5, INK, after=4, line=1.32)
        if notes:
            if isinstance(notes, list):
                add_bullets(doc, notes)
            else:
                add_paragraph_text(doc, notes)
        if flow:
            add_flow(doc, flow)
        if image:
            add_discussion_image(doc, image)
        if highlight:
            add_highlight_box(doc, highlight)


def _table(doc, headers, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.allow_autofit = False
    set_table_borders(tbl, HAIRLINE, 4)
    for c, w in enumerate(widths):
        tbl.columns[c].width = w
    hdr = tbl.rows[0]
    for c, h in enumerate(headers):
        cell = hdr.cells[c]
        first_cell_clear(cell)
        cell.width = widths[c]
        shade_cell(cell, INK_HEX)
        set_cell_margins(cell, 60, 60, 100, 100)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(h.upper()), BODY, 8.5, WHITE, bold=True, spacing=6)
    return tbl


def _fill_row(tbl, values, widths, zebra=False):
    row = tbl.add_row()
    for c, val in enumerate(values):
        cell = row.cells[c]
        cell.width = widths[c]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_margins(cell, 90, 90, 110, 110)
        if zebra:
            shade_cell(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        style_run(p.add_run("" if val is None else str(val)), BODY, 9.5, INK)
    return row


def add_action_items(doc, items):
    widths = [Cm(1.0), Cm(8.2), Cm(3.0), Cm(2.6), Cm(2.2)]
    tbl = _table(doc, ["#", "Action", "Owner", "Due", "Status"], widths)
    for i, it in enumerate(items, 1):
        it = it if isinstance(it, dict) else {"action": str(it)}
        _fill_row(tbl, [
            i,
            it.get("action") or it.get("task") or "",
            it.get("owner") or it.get("responsible") or "",
            it.get("due") or it.get("due_date") or it.get("deadline") or "",
            it.get("status") or "Open",
        ], widths, zebra=(i % 2 == 0))


def add_risks(doc, items):
    widths = [Cm(1.0), Cm(5.2), Cm(2.8), Cm(5.6), Cm(2.4)]
    tbl = _table(doc, ["#", "Risk / Blocker", "Impact", "Workaround / Mitigation", "Owner"], widths)
    for i, it in enumerate(items, 1):
        it = it if isinstance(it, dict) else {"risk": str(it)}
        _fill_row(tbl, [
            i,
            it.get("risk") or it.get("blocker") or it.get("issue") or "",
            it.get("impact") or it.get("severity") or "",
            it.get("workaround") or it.get("mitigation") or it.get("resolution") or "",
            it.get("owner") or it.get("responsible") or "",
        ], widths, zebra=(i % 2 == 0))


def add_participants(doc, people):
    widths = [Cm(0.9), Cm(5.6), Cm(5.5), Cm(5.0)]
    tbl = _table(doc, ["#", "Name", "Role", "Organization"], widths)
    for i, p in enumerate(people, 1):
        p = p if isinstance(p, dict) else {"name": str(p)}
        _fill_row(tbl, [
            i,
            p.get("name") or "",
            p.get("role") or p.get("title") or "",
            p.get("organization") or p.get("org") or p.get("company") or "",
        ], widths, zebra=(i % 2 == 0))


def _row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def add_approval_section(doc, approvers, lang="en"):
    """Approval heading + signature blocks, kept together as one unit.

    The whole block uses "keep with next" + "can't split row" so it flows
    naturally at the bottom of the page, and only moves to the next page as a
    whole if it doesn't fit — it is never split across pages, and we don't force
    a page break of its own.
    """
    intro_txt = ("Dokumen ini disetujui oleh:" if lang == "id"
                 else "This document is approved by:")
    date_lbl = "Tanggal: ____________" if lang == "id" else "Date: ____________"

    # same uniform section header as every other section
    add_section_header(doc, "Approval", keep_next=True)

    intro = add_para(doc, intro_txt, BODY, 9.5, MUTED, after=10)
    intro.paragraph_format.keep_with_next = True

    per_row = 2
    col_w = Cm(8.5)
    for start in range(0, len(approvers), per_row):
        group = approvers[start:start + per_row]
        tbl = doc.add_table(rows=1, cols=per_row)
        tbl.allow_autofit = False
        no_table_borders(tbl)
        _row_cant_split(tbl.rows[0])
        for c in range(per_row):
            tbl.columns[c].width = col_w
        for c in range(per_row):
            cell = tbl.cell(0, c)
            first_cell_clear(cell)
            cell.width = col_w
            set_cell_margins(cell, 60, 60, 0, 220)
            if c >= len(group):
                cell.add_paragraph()
                continue
            a = group[c] if isinstance(group[c], dict) else {"name": str(group[c])}
            sp = cell.add_paragraph()
            sp.paragraph_format.space_after = Pt(0)
            sp.paragraph_format.line_spacing = Pt(46)
            line = cell.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            cell_para_bottom_border(line, INK_HEX, 6)
            np = cell.add_paragraph()
            np.paragraph_format.space_after = Pt(0)
            style_run(np.add_run(a.get("name") or "(__________________)"),
                      BODY, 10, INK, bold=True)
            role = a.get("role") or a.get("title") or ""
            org = a.get("organization") or a.get("org") or a.get("company") or ""
            detail = ", ".join([x for x in (role, org) if x])
            if detail:
                dp = cell.add_paragraph()
                dp.paragraph_format.space_after = Pt(0)
                style_run(dp.add_run(detail), BODY, 9, MUTED)
            datep = cell.add_paragraph()
            datep.paragraph_format.space_before = Pt(2)
            datep.paragraph_format.space_after = Pt(0)
            style_run(datep.add_run(date_lbl), BODY, 8.5, MUTED)


def cell_para_bottom_border(par, color=INK_HEX, size=6):
    """Bottom border on a paragraph (used as a signature line)."""
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)
    bottom.set(qn("w:space"), "1")
    pbdr.append(bottom)
    pPr.append(pbdr)


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
def add_footer(doc, data):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    # Detach from the built-in "Footer" style, which carries its own
    # center/right default tab stops that would override ours.
    try:
        p.style = doc.styles["Normal"]
    except KeyError:
        pass
    p.paragraph_format.space_before = Pt(2)
    # hairline above footer
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:color"), HAIRLINE)
    top.set(qn("w:space"), "6")
    pbdr.append(top)
    pPr.append(pbdr)

    style_run(p.add_run(TAGLINE), BODY, 7.5, MUTED)
    p.add_run("\t")
    conf = data.get("confidentiality")
    if conf is None:
        conf = "Confidential" if (data.get("meeting_type", "").lower() == "external") else "Internal"
    style_run(p.add_run(f"{conf}   |   Page "), BODY, 7.5, MUTED)
    _page_field(p)
    style_run(p.add_run(" of "), BODY, 7.5, MUTED)
    _page_field(p, "NUMPAGES")

    # right-align the second half via tab stop at content edge
    from docx.enum.text import WD_TAB_ALIGNMENT
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)


def _page_field(paragraph, field="PAGE"):
    run = paragraph.add_run()
    style_run(run, BODY, 7.5, MUTED)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------
def _looks_indonesian(data):
    """Cheap fallback when `language` isn't set: scan the body for very common
    Indonesian words. Only used to pick the approval label language."""
    text = " ".join([
        str(data.get("executive_summary", "")),
        " ".join(str(x) for x in data.get("key_points", []) if isinstance(x, str)),
        " ".join(str(x) for x in data.get("next_steps", []) if isinstance(x, str)),
    ]).lower()
    hits = sum(text.count(w) for w in
               (" yang ", " dan ", " untuk ", " dengan ", " pada ", " tidak ", " adalah ", " ini "))
    return hits >= 2


def build(data, out_basename):
    doc = Document()
    doc.core_properties.title = data.get("meeting_title", "Minutes of Meeting")
    doc.core_properties.author = data.get("prepared_by", "BUBU")

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    add_header_band(doc, data)
    add_meta_block(doc, data)

    def section(label, key, renderer):
        val = data.get(key)
        if not val:
            return
        add_section_header(doc, label)
        renderer(val)

    # Participants sit near the top so the reader knows who was in the room
    # before reading what was discussed.
    section("Participants", "participants",
            lambda v: add_participants(doc, v))
    section("Executive Summary", "executive_summary",
            lambda v: add_paragraph_text(doc, v))
    section("Agenda", "agenda",
            lambda v: add_bullets(doc, v, numbered=True))
    section("Key Points & Decisions", "key_points",
            lambda v: add_bullets(doc, v))
    section("Discussion", "discussion",
            lambda v: add_discussion(doc, v))
    section("Action Items", "action_items",
            lambda v: add_action_items(doc, v))
    section("Risks, Blockers & Workarounds", "risks",
            lambda v: add_risks(doc, v))
    section("Next Steps", "next_steps",
            lambda v: add_bullets(doc, v))

    # Approval closes the document. It flows naturally at the bottom of the page;
    # "keep together" means it only moves to the next page (as a whole) if it
    # doesn't fit — it never starts a page of its own and is never split.
    if data.get("approvals"):
        lang = (data.get("language") or "").strip().lower()
        if lang not in ("id", "en"):
            lang = "id" if _looks_indonesian(data) else "en"
        add_approval_section(doc, data["approvals"], lang)

    add_footer(doc, data)

    docx_path = out_basename + ".docx"
    doc.save(docx_path)
    try:
        embed_font_in_docx(docx_path)
    except Exception as e:  # noqa: BLE001 — never let embedding break generation
        sys.stderr.write(f"[warn] font embedding skipped: {e}\n")
    return docx_path


def embed_font_in_docx(docx_path, font_name=DISPLAY, ttf_name="BebasNeue-Regular.ttf"):
    """Embed the display font directly inside the .docx so headings render in
    Bebas Neue on ANY computer — even one that doesn't have the font installed.

    A .docx only references fonts by name; if the reader's machine lacks the
    font, Word/Pages silently substitutes another. Embedding the actual font
    bytes (obfuscated per the OOXML spec, the same mechanism Word uses) makes the
    document self-sufficient. This is what the original BUBU template did.
    """
    import uuid
    import zipfile

    font_path = os.path.join(FONT_DIR, ttf_name)
    if not os.path.exists(font_path):
        return

    # --- obfuscate the first 32 bytes with a GUID key (OOXML embedded font) ---
    guid_hex = uuid.uuid4().hex.upper()  # 32 hex chars
    guid_str = (f"{{{guid_hex[0:8]}-{guid_hex[8:12]}-{guid_hex[12:16]}-"
                f"{guid_hex[16:20]}-{guid_hex[20:32]}}}")
    key = bytes.fromhex(guid_hex)        # 16 bytes
    data = bytearray(open(font_path, "rb").read())
    for i in range(32):
        data[i] ^= key[15 - (i % 16)]
    obf = bytes(data)

    odttf_part = "word/fonts/font_bebasneue.odttf"
    rel_id = "rIdBebasNeue"

    zin = zipfile.ZipFile(docx_path, "r")
    items = {n: zin.read(n) for n in zin.namelist()}
    zin.close()

    # 1) [Content_Types].xml — register the .odttf extension
    ct = items["[Content_Types].xml"].decode("utf-8")
    if "obfuscatedFont" not in ct:
        ct = ct.replace(
            "</Types>",
            '<Default Extension="odttf" '
            'ContentType="application/vnd.openxmlformats-officedocument.obfuscatedFont"/></Types>',
        )
    items["[Content_Types].xml"] = ct.encode("utf-8")

    # 2) fontTable.xml — declare the font as embedded
    ft = items["word/fontTable.xml"].decode("utf-8")
    font_entry = (
        f'<w:font w:name="{font_name}">'
        f'<w:embedRegular r:id="{rel_id}" w:fontKey="{guid_str}" w:subsetted="0"/>'
        f"</w:font>"
    )
    ft = ft.replace("</w:fonts>", font_entry + "</w:fonts>")
    items["word/fontTable.xml"] = ft.encode("utf-8")

    # 3) fontTable rels — point the relationship at the embedded part
    rels_name = "word/_rels/fontTable.xml.rels"
    rel = ('<Relationship '
           f'Id="{rel_id}" '
           'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" '
           'Target="fonts/font_bebasneue.odttf"/>')
    if rels_name in items:
        r = items[rels_name].decode("utf-8")
        r = r.replace("</Relationships>", rel + "</Relationships>")
    else:
        r = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
             '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
             + rel + "</Relationships>")
    items[rels_name] = r.encode("utf-8")

    # 4) settings.xml — turn on font embedding
    st = items["word/settings.xml"].decode("utf-8")
    if "embedTrueTypeFonts" not in st:
        st = st.replace("<w:settings ", "<w:settings ", 1)
        # insert right after the opening <w:settings ...> tag
        idx = st.find(">", st.find("<w:settings")) + 1
        st = (st[:idx]
              + '<w:embedTrueTypeFonts/><w:embedSystemFonts/><w:saveSubsetFonts w:val="0"/>'
              + st[idx:])
    items["word/settings.xml"] = st.encode("utf-8")

    # 5) add the obfuscated font part
    items[odttf_part] = obf

    # rewrite the archive
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, blob in items.items():
            zout.writestr(name, blob)
    os.replace(tmp, docx_path)


def ensure_fonts():
    """Register the bundled Bebas Neue with fontconfig so the brand typography
    survives PDF conversion on a fresh machine.

    Each session/sandbox starts with a clean font set, so we cannot assume the
    display font is installed. We copy the bundled TTF(s) into the user font dir
    and refresh the cache. Safe to run repeatedly; failures degrade gracefully
    (the .docx still names the font correctly for users who have it installed).
    """
    if not os.path.isdir(FONT_DIR):
        return
    home = os.path.expanduser("~")
    targets = [os.path.join(home, ".fonts"),
               os.path.join(home, ".local", "share", "fonts")]
    installed = False
    for dest in targets:
        try:
            os.makedirs(dest, exist_ok=True)
            for fn in os.listdir(FONT_DIR):
                if fn.lower().endswith((".ttf", ".otf")):
                    src = os.path.join(FONT_DIR, fn)
                    dst = os.path.join(dest, fn)
                    if not os.path.exists(dst):
                        with open(src, "rb") as a, open(dst, "wb") as b:
                            b.write(a.read())
            installed = True
            break
        except Exception:  # noqa: BLE001
            continue
    if installed:
        try:
            subprocess.run(["fc-cache", "-f"], capture_output=True, timeout=60)
        except Exception:  # noqa: BLE001
            pass


def to_pdf(docx_path):
    """Convert to PDF inside an isolated temp dir, then move only the finished
    PDF next to the .docx. LibreOffice scatters lock files (.~lock.*#) and temp
    files in its working/output dir; doing the work in a throwaway dir keeps the
    user's folder clean — just the .docx and .pdf, nothing else.
    """
    ensure_fonts()
    import shutil
    import tempfile

    final_pdf = os.path.splitext(docx_path)[0] + ".pdf"
    try:
        with tempfile.TemporaryDirectory() as td:
            tmp_docx = os.path.join(td, os.path.basename(docx_path))
            shutil.copy2(docx_path, tmp_docx)
            profile = "file://" + os.path.join(td, "loprofile")
            subprocess.run(
                ["soffice", f"-env:UserInstallation={profile}", "--headless",
                 "--convert-to", "pdf", "--outdir", td, tmp_docx],
                check=True, capture_output=True, timeout=120,
            )
            tmp_pdf = os.path.splitext(tmp_docx)[0] + ".pdf"
            if os.path.exists(tmp_pdf):
                shutil.move(tmp_pdf, final_pdf)
                return final_pdf
        return None
    except Exception as e:  # noqa: BLE001
        sys.stderr.write(f"[warn] PDF conversion skipped: {e}\n")
        return None


def _safe_name(text, fallback="Meeting"):
    """Make a string safe for a file/folder name (keep spaces, drop separators)."""
    text = (text or "").strip()
    if not text:
        text = fallback
    bad = '/\\:*?"<>|'
    cleaned = "".join("-" if c in bad else c for c in text)
    cleaned = " ".join(cleaned.split())  # collapse whitespace
    return cleaned[:80].strip(" .-") or fallback


def compute_output_basename(data, root):
    """Build <root>/MOM/<Company|Internal>/<Name> MOM - <date> (no extension).

    External meetings file under the client/company name; internal meetings file
    under an "Internal" folder. The filename leads with the company (external) or
    the project/meeting title (internal) so the folder reads cleanly at a glance.
    """
    is_internal = (data.get("meeting_type") or "").strip().lower() == "internal"

    # `company` lets the user override the long legal name with a short tidy one.
    company = data.get("company") or data.get("client") or data.get("project")

    if is_internal:
        subfolder = "Internal"
        name_part = data.get("company") or data.get("project") or data.get("meeting_title") or "Internal"
    else:
        subfolder = _safe_name(company, "External")
        name_part = _safe_name(company, "Meeting")

    date = data.get("date_iso") or datetime.now().strftime("%Y-%m-%d")
    filename = f"{_safe_name(name_part)} MOM - {date}"

    folder = os.path.join(root, "MOM", subfolder)
    os.makedirs(folder, exist_ok=True)
    return os.path.join(folder, filename)


def main():
    args = sys.argv[1:]
    root = None
    if "--root" in args:
        i = args.index("--root")
        try:
            root = args[i + 1]
        except IndexError:
            sys.exit("--root requires a directory path")
        del args[i:i + 2]

    if not args:
        sys.exit("Usage: python3 generate_mom.py mom_data.json [output_basename] "
                 "[--root BUBU_FOLDER]")

    data_path = args[0]
    with open(data_path, "r", encoding="utf-8") as fh:
        data = json.load(fh)

    if len(args) >= 2:
        # explicit output path always wins
        out_basename = os.path.splitext(args[1])[0]
    elif root:
        # auto-file into BUBU/MOM/<Company|Internal>/<Company> MOM - <date>
        out_basename = compute_output_basename(data, root)
    else:
        slug = (data.get("project") or data.get("meeting_title") or "Meeting")
        slug = "".join(c if c.isalnum() else "_" for c in slug).strip("_")[:40]
        date = data.get("date_iso") or datetime.now().strftime("%Y-%m-%d")
        out_basename = f"MOM_{slug}_{date}"

    docx_path = build(data, out_basename)
    pdf_path = to_pdf(docx_path)
    print(f"DOCX: {docx_path}")
    print(f"PDF:  {pdf_path or 'not generated (LibreOffice unavailable)'}")


if __name__ == "__main__":
    main()
