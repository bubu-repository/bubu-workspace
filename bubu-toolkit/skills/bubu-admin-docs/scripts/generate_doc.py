#!/usr/bin/env python3
"""
BUBU Administrative Document generator.

One engine for every BUBU administrative document — official letters, memos,
proposals, quotations, invoices, SOPs, policies, berita acara, forms, and more —
all sharing the exact same BUBU visual system as the BUBU MOM skill (orange
#FF5900, Bebas Neue headings, Arial body, logo header band, footer, uniform
spacing, embedded font).

Content is a flexible list of "blocks", so the same template covers any admin
document. Output is .docx with Bebas Neue embedded, so it looks right on any
computer.

Usage:
    python3 generate_doc.py doc_data.json [output_basename] [--root BUBU_FOLDER]

See references/data_schema.md for the JSON schema.
"""

import json
import os
import sys
import zipfile
import uuid
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# BUBU brand system  (identical to the BUBU MOM skill)
# ---------------------------------------------------------------------------
ORANGE = RGBColor(0xFF, 0x59, 0x00)
INK = RGBColor(0x14, 0x14, 0x16)
MUTED = RGBColor(0x70, 0x70, 0x74)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HAIRLINE = "E3E3E3"
LIGHT = "F5F4F2"
ORANGE_HEX = "FF5900"
INK_HEX = "141416"

DISPLAY = "Bebas Neue"
BODY = "Arial"

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.normpath(os.path.join(HERE, "..", "assets"))
LOGO = os.path.join(ASSETS, "bubu_logo.png")
FONT_DIR = os.path.join(ASSETS, "fonts")

TAGLINE = "BUBU · 30 years turning culture into brands, products & experiences"

CONTENT_W = Cm(17.0)  # A4 (21cm) minus 2cm margins each side

# Uniform spacing around section headings (points)
SECTION_GAP_BEFORE = 13
SECTION_GAP_AFTER = 6


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def _shade(element, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def shade_cell(cell, fill_hex):
    _shade(cell._tc.get_or_add_tcPr(), fill_hex)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                     ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def no_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def set_table_borders(table, color=HAIRLINE, size=4):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:color"), color)
        borders.append(e)
    table._tbl.tblPr.append(borders)


def cell_para_bottom_border(par, color=INK_HEX, size=6):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)
    bottom.set(qn("w:space"), "1")
    pbdr.append(bottom)
    pPr.append(pbdr)


def cell_borders(cell, color=HAIRLINE, size=4, left_color=None, left_size=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        if edge == "left" and left_color:
            e.set(qn("w:sz"), str(left_size or size))
            e.set(qn("w:color"), left_color)
        else:
            e.set(qn("w:sz"), str(size))
            e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def style_run(run, font=BODY, size=10.5, color=INK, bold=False, caps=False,
              italic=False, spacing=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.all_caps = caps
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)
    if spacing is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rpr.append(sp)
    return run


def add_para(container, text="", font=BODY, size=10.5, color=INK, bold=False,
             caps=False, align=None, before=0, after=4, line=1.12, spacing=None,
             italic=False):
    par = container.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line
    if align is not None:
        par.alignment = align
    if text != "":
        style_run(par.add_run(text), font, size, color, bold, caps, italic, spacing)
    return par


def first_cell_clear(cell):
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)


def _row_cant_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


# ---------------------------------------------------------------------------
# Header / meta / footer
# ---------------------------------------------------------------------------
def add_header_band(doc, data):
    """Logo (left) + optional tag (right) + orange band with the document title."""
    top = doc.add_table(rows=1, cols=2)
    top.allow_autofit = False
    no_table_borders(top)
    top.columns[0].width = Cm(11.0)
    top.columns[1].width = Cm(6.0)

    logo_cell = top.cell(0, 0)
    first_cell_clear(logo_cell)
    set_cell_margins(logo_cell, 0, 0, 0, 0)
    lp = logo_cell.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO):
        lp.add_run().add_picture(LOGO, height=Cm(0.95))

    tag_cell = top.cell(0, 1)
    first_cell_clear(tag_cell)
    tag_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(tag_cell, 0, 0, 0, 0)
    tag = (data.get("tag") or data.get("confidentiality") or "").strip()
    tp = tag_cell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tp.paragraph_format.space_after = Pt(0)
    if tag:
        style_run(tp.add_run(tag.upper()), DISPLAY, 13, ORANGE, spacing=20)

    spacer = add_para(doc, "", before=0, after=0)
    spacer.paragraph_format.line_spacing = Pt(10)

    band = doc.add_table(rows=1, cols=1)
    band.allow_autofit = False
    no_table_borders(band)
    band.columns[0].width = CONTENT_W
    bc = band.cell(0, 0)
    first_cell_clear(bc)
    bc.width = CONTENT_W
    shade_cell(bc, ORANGE_HEX)
    set_cell_margins(bc, 150, 150, 220, 160)

    title = data.get("doc_type") or data.get("title") or "Document"
    tpar = bc.add_paragraph()
    tpar.paragraph_format.space_after = Pt(0)
    tpar.paragraph_format.line_spacing = 0.95
    style_run(tpar.add_run(title.upper()), DISPLAY, 30, WHITE, spacing=20)

    # optional one-line subtitle under the title (e.g. the document's specific topic)
    sub = data.get("subtitle") or ""
    if sub:
        sp2 = bc.add_paragraph()
        sp2.paragraph_format.space_before = Pt(2)
        sp2.paragraph_format.space_after = Pt(0)
        style_run(sp2.add_run(sub), BODY, 11.5, WHITE, bold=True)

    add_para(doc, "", after=6)


def add_meta_block(doc, meta):
    """Two-column key/value grid. `meta` is a list of {label, value}."""
    pairs = []
    for item in meta or []:
        if isinstance(item, dict):
            label = item.get("label") or ""
            value = item.get("value") or ""
        else:
            continue
        if label or value:
            pairs.append((str(label), str(value)))
    if not pairs:
        return

    rows = (len(pairs) + 1) // 2
    tbl = doc.add_table(rows=rows, cols=4)
    tbl.allow_autofit = False
    no_table_borders(tbl)
    widths = [Cm(3.4), Cm(5.1), Cm(3.4), Cm(5.1)]
    for r in range(rows):
        for c in range(4):
            tbl.cell(r, c).width = widths[c]

    for i, (label, value) in enumerate(pairs):
        r = i // 2
        base = (i % 2) * 2
        lc = tbl.cell(r, base)
        vc = tbl.cell(r, base + 1)
        first_cell_clear(lc)
        first_cell_clear(vc)
        set_cell_margins(lc, 70, 70, 40, 80)
        set_cell_margins(vc, 70, 70, 40, 120)
        shade_cell(lc, LIGHT)
        shade_cell(vc, LIGHT)
        lp = lc.add_paragraph()
        lp.paragraph_format.space_after = Pt(0)
        style_run(lp.add_run(label.upper()), BODY, 8, MUTED, bold=True, spacing=8)
        vp = vc.add_paragraph()
        vp.paragraph_format.space_after = Pt(0)
        style_run(vp.add_run(value), BODY, 10, INK)

    if len(pairs) % 2 == 1:
        for c in (2, 3):
            shade_cell(tbl.cell(rows - 1, c), LIGHT)


def add_footer(doc, data):
    from docx.enum.text import WD_TAB_ALIGNMENT
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    try:
        p.style = doc.styles["Normal"]
    except KeyError:
        pass
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:color"), HAIRLINE)
    top.set(qn("w:space"), "6")
    pbdr.append(top)
    pPr.append(pbdr)

    style_run(p.add_run(TAGLINE), BODY, 7.5, MUTED)
    p.add_run("\t")
    conf = data.get("confidentiality") or "Internal"
    style_run(p.add_run(f"{conf}   |   Page "), BODY, 7.5, MUTED)
    _page_field(p)
    style_run(p.add_run(" of "), BODY, 7.5, MUTED)
    _page_field(p, "NUMPAGES")
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)


def _page_field(paragraph, field="PAGE"):
    run = paragraph.add_run()
    style_run(run, BODY, 7.5, MUTED)
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1)
    run._r.append(instr)
    run._r.append(f2)


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------
def add_section_header(doc, label, keep_next=True):
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.space_before = Pt(SECTION_GAP_BEFORE)
    pf.space_after = Pt(SECTION_GAP_AFTER)
    pf.keep_with_next = keep_next
    cell_para_bottom_border(par, ORANGE_HEX, 16)
    style_run(par.add_run("■  "), DISPLAY, 14, ORANGE)
    style_run(par.add_run(str(label).upper()), DISPLAY, 15, INK, spacing=14)
    return par


def add_paragraphs(doc, text):
    blocks = [b.strip() for b in str(text).split("\n") if b.strip()]
    for i, b in enumerate(blocks):
        add_para(doc, b, BODY, 10.5, INK, after=(0 if i == len(blocks) - 1 else 8), line=1.4)


def add_list(doc, items, numbered=False):
    n = len(items)
    for i, item in enumerate(items, 1):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.7)
        par.paragraph_format.space_after = Pt(0 if i == n else 5)
        par.paragraph_format.line_spacing = 1.32
        marker = f"{i}." if numbered else "•"
        style_run(par.add_run(f"{marker}  "), BODY, 10.5, ORANGE, bold=True)
        style_run(par.add_run(str(item)), BODY, 10.5, INK)


def add_highlight(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.allow_autofit = False
    no_table_borders(tbl)
    tbl.columns[0].width = CONTENT_W
    cell = tbl.cell(0, 0)
    first_cell_clear(cell)
    cell.width = CONTENT_W
    shade_cell(cell, LIGHT)
    set_cell_margins(cell, 130, 130, 200, 160)
    cell_borders(cell, color="DEDDD8", size=4, left_color=ORANGE_HEX, left_size=26)
    blocks = [b.strip() for b in str(text).split("\n") if b.strip()]
    for i, b in enumerate(blocks):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0 if i == len(blocks) - 1 else 4)
        p.paragraph_format.line_spacing = 1.3
        style_run(p.add_run(b), BODY, 10.5, INK, italic=True)


def _coerce_str(v):
    return "" if v is None else str(v)


def add_data_table(doc, block):
    """Generic table with optional right-aligned columns and a totals section.

    block = {
      "columns": ["No","Description","Qty","Price","Amount"],
      "rows": [[...], ...],
      "align_right": [2,3,4],          # column indices rendered right-aligned (optional)
      "totals": [["Subtotal","Rp ..."],["Total","Rp ..."]]   # optional; last row bold
    }
    """
    columns = block.get("columns") or []
    rows = block.get("rows") or []
    if not columns:
        return
    ncol = len(columns)
    align_right = set(block.get("align_right") or [])

    # distribute widths: first column narrow if it looks like a number column
    total_w = 17.0
    widths_cm = [total_w / ncol] * ncol
    if columns and columns[0].strip().lower() in ("no", "no.", "#"):
        widths_cm[0] = 1.0
        rest = (total_w - 1.0) / (ncol - 1) if ncol > 1 else total_w
        widths_cm = [1.0] + [rest] * (ncol - 1)
    widths = [Cm(w) for w in widths_cm]

    tbl = doc.add_table(rows=1, cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.allow_autofit = False
    set_table_borders(tbl, HAIRLINE, 4)
    for c in range(ncol):
        tbl.columns[c].width = widths[c]

    hdr = tbl.rows[0]
    for c, h in enumerate(columns):
        cell = hdr.cells[c]
        first_cell_clear(cell)
        cell.width = widths[c]
        shade_cell(cell, INK_HEX)
        set_cell_margins(cell, 70, 70, 110, 110)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        if c in align_right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_run(p.add_run(str(h).upper()), BODY, 8.5, WHITE, bold=True, spacing=6)

    for ri, row in enumerate(rows):
        tr = tbl.add_row()
        for c in range(ncol):
            cell = tr.cells[c]
            cell.width = widths[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell, 80, 80, 110, 110)
            if ri % 2 == 1:
                shade_cell(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            if c in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            val = row[c] if c < len(row) else ""
            style_run(p.add_run(_coerce_str(val)), BODY, 9.5, INK)

    # totals block: a small right-aligned two-column summary under the table
    totals = block.get("totals") or []
    if totals:
        add_para(doc, "", after=2)
        tt = doc.add_table(rows=0, cols=2)
        tt.allow_autofit = False
        no_table_borders(tt)
        lw, rw = Cm(12.5), Cm(4.5)
        tt.columns[0].width = lw
        tt.columns[1].width = rw
        for i, pair in enumerate(totals):
            label = pair[0] if len(pair) > 0 else ""
            value = pair[1] if len(pair) > 1 else ""
            is_last = (i == len(totals) - 1)
            r = tt.add_row()
            lc, rc = r.cells[0], r.cells[1]
            lc.width, rc.width = lw, rw
            set_cell_margins(lc, 50, 50, 60, 120)
            set_cell_margins(rc, 50, 50, 60, 120)
            if is_last:
                shade_cell(lc, LIGHT)
                shade_cell(rc, LIGHT)
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lp.paragraph_format.space_after = Pt(0)
            style_run(lp.add_run(str(label)), BODY, 10, INK, bold=is_last)
            rp = rc.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rp.paragraph_format.space_after = Pt(0)
            style_run(rp.add_run(str(value)), BODY, 10, (ORANGE if is_last else INK), bold=is_last)


def add_signatures(doc, block, lang="en"):
    approvers = block.get("approvers") or block.get("signatories") or []
    if not approvers:
        return
    date_lbl = "Tanggal: ____________" if lang == "id" else "Date: ____________"
    intro = block.get("intro")
    if intro:
        p = add_para(doc, intro, BODY, 9.5, MUTED, after=10)
        p.paragraph_format.keep_with_next = True

    per_row = 2
    col_w = Cm(8.5)
    for start in range(0, len(approvers), per_row):
        group = approvers[start:start + per_row]
        tbl = doc.add_table(rows=1, cols=per_row)
        tbl.allow_autofit = False
        no_table_borders(tbl)
        _row_cant_split(tbl.rows[0])
        for c in range(per_row):
            tbl.columns[c].width = col_w
        for c in range(per_row):
            cell = tbl.cell(0, c)
            first_cell_clear(cell)
            cell.width = col_w
            set_cell_margins(cell, 60, 60, 0, 220)
            if c >= len(group):
                cell.add_paragraph()
                continue
            a = group[c] if isinstance(group[c], dict) else {"name": str(group[c])}
            sp = cell.add_paragraph()
            sp.paragraph_format.space_after = Pt(0)
            sp.paragraph_format.line_spacing = Pt(46)
            line = cell.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            cell_para_bottom_border(line, INK_HEX, 6)
            np = cell.add_paragraph()
            np.paragraph_format.space_after = Pt(0)
            style_run(np.add_run(a.get("name") or "(__________________)"), BODY, 10, INK, bold=True)
            role = a.get("role") or a.get("title") or ""
            org = a.get("organization") or a.get("org") or a.get("company") or ""
            detail = ", ".join([x for x in (role, org) if x])
            if detail:
                dp = cell.add_paragraph()
                dp.paragraph_format.space_after = Pt(0)
                style_run(dp.add_run(detail), BODY, 9, MUTED)
            datep = cell.add_paragraph()
            datep.paragraph_format.space_before = Pt(2)
            datep.paragraph_format.space_after = Pt(0)
            style_run(datep.add_run(date_lbl), BODY, 8.5, MUTED)


def render_blocks(doc, blocks, lang="en"):
    for blk in blocks or []:
        if isinstance(blk, str):
            blk = {"type": "paragraph", "text": blk}
        btype = (blk.get("type") or "paragraph").lower()
        if btype in ("heading", "section", "header"):
            add_section_header(doc, blk.get("text") or blk.get("title") or "")
        elif btype in ("paragraph", "para", "text"):
            add_paragraphs(doc, blk.get("text") or "")
        elif btype in ("bullets", "list", "bullet"):
            add_list(doc, blk.get("items") or [], numbered=False)
        elif btype in ("numbered", "ordered", "steps"):
            add_list(doc, blk.get("items") or [], numbered=True)
        elif btype in ("table", "items", "invoice"):
            add_data_table(doc, blk)
        elif btype in ("highlight", "callout", "note"):
            add_highlight(doc, blk.get("text") or "")
        elif btype in ("signatures", "approval", "signoff"):
            add_section_header(doc, blk.get("label") or ("Persetujuan" if lang == "id" else "Approval"))
            add_signatures(doc, blk, lang)
        elif btype == "spacer":
            add_para(doc, "", after=blk.get("size", 6))
        else:
            add_paragraphs(doc, blk.get("text") or "")


# ---------------------------------------------------------------------------
# Font embedding (so Bebas Neue renders on any computer)
# ---------------------------------------------------------------------------
def embed_font_in_docx(docx_path, font_name=DISPLAY, ttf_name="BebasNeue-Regular.ttf"):
    font_path = os.path.join(FONT_DIR, ttf_name)
    if not os.path.exists(font_path):
        return
    guid_hex = uuid.uuid4().hex.upper()
    guid_str = (f"{{{guid_hex[0:8]}-{guid_hex[8:12]}-{guid_hex[12:16]}-"
                f"{guid_hex[16:20]}-{guid_hex[20:32]}}}")
    key = bytes.fromhex(guid_hex)
    data = bytearray(open(font_path, "rb").read())
    for i in range(32):
        data[i] ^= key[15 - (i % 16)]
    obf = bytes(data)

    odttf = "word/fonts/font_bebasneue.odttf"
    rel_id = "rIdBebasNeue"
    zin = zipfile.ZipFile(docx_path, "r")
    items = {n: zin.read(n) for n in zin.namelist()}
    zin.close()

    ct = items["[Content_Types].xml"].decode("utf-8")
    if "obfuscatedFont" not in ct:
        ct = ct.replace("</Types>",
                        '<Default Extension="odttf" '
                        'ContentType="application/vnd.openxmlformats-officedocument.obfuscatedFont"/></Types>')
    items["[Content_Types].xml"] = ct.encode("utf-8")

    if "word/fontTable.xml" in items:
        ft = items["word/fontTable.xml"].decode("utf-8")
    else:
        ft = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
              '<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
              'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"></w:fonts>')
    entry = (f'<w:font w:name="{font_name}">'
             f'<w:embedRegular r:id="{rel_id}" w:fontKey="{guid_str}" w:subsetted="0"/></w:font>')
    ft = ft.replace("</w:fonts>", entry + "</w:fonts>")
    items["word/fontTable.xml"] = ft.encode("utf-8")

    rels_name = "word/_rels/fontTable.xml.rels"
    rel = ('<Relationship '
           f'Id="{rel_id}" '
           'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/font" '
           'Target="fonts/font_bebasneue.odttf"/>')
    if rels_name in items:
        r = items[rels_name].decode("utf-8").replace("</Relationships>", rel + "</Relationships>")
    else:
        r = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
             '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
             + rel + "</Relationships>")
    items[rels_name] = r.encode("utf-8")

    st = items["word/settings.xml"].decode("utf-8")
    if "embedTrueTypeFonts" not in st:
        idx = st.find(">", st.find("<w:settings")) + 1
        st = st[:idx] + '<w:embedTrueTypeFonts/><w:embedSystemFonts/><w:saveSubsetFonts w:val="0"/>' + st[idx:]
    items["word/settings.xml"] = st.encode("utf-8")

    # ensure fontTable part is declared in content types
    ct = items["[Content_Types].xml"].decode("utf-8")
    if "fontTable+xml" not in ct:
        ct = ct.replace("</Types>",
                        '<Override PartName="/word/fontTable.xml" '
                        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.fontTable+xml"/></Types>')
        items["[Content_Types].xml"] = ct.encode("utf-8")
    # ensure fontTable is referenced from document rels
    drels = "word/_rels/document.xml.rels"
    if drels in items:
        dr = items[drels].decode("utf-8")
        if "fontTable.xml" not in dr:
            dr = dr.replace("</Relationships>",
                            '<Relationship Id="rIdFontTable" '
                            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable" '
                            'Target="fontTable.xml"/></Relationships>')
            items[drels] = dr.encode("utf-8")

    items[odttf] = obf
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, blob in items.items():
            zout.writestr(name, blob)
    os.replace(tmp, docx_path)


# ---------------------------------------------------------------------------
# Assembly
# ---------------------------------------------------------------------------
def build(data, out_basename):
    doc = Document()
    doc.core_properties.title = data.get("doc_type") or data.get("title") or "Document"
    doc.core_properties.author = data.get("prepared_by") or "BUBU"

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    lang = (data.get("language") or "").strip().lower()
    if lang not in ("id", "en"):
        lang = "en"

    add_header_band(doc, data)
    add_meta_block(doc, data.get("meta"))
    if data.get("meta"):
        add_para(doc, "", after=0)
    render_blocks(doc, data.get("body"), lang)
    add_footer(doc, data)

    docx_path = out_basename + ".docx"
    doc.save(docx_path)
    try:
        embed_font_in_docx(docx_path)
    except Exception as e:  # noqa: BLE001
        sys.stderr.write(f"[warn] font embedding skipped: {e}\n")
    return docx_path


def _safe_name(text, fallback="Document"):
    text = (text or "").strip() or fallback
    bad = '/\\:*?"<>|'
    cleaned = "".join("-" if c in bad else c for c in text)
    return " ".join(cleaned.split())[:80].strip(" .-") or fallback


def compute_output_basename(data, root):
    doc_type = _safe_name(data.get("doc_type") or "Document")
    name_part = _safe_name(data.get("title") or data.get("doc_type") or "Document")
    date = data.get("date_iso") or datetime.now().strftime("%Y-%m-%d")
    folder = os.path.join(root, "Administration", doc_type)
    os.makedirs(folder, exist_ok=True)
    return os.path.join(folder, f"{name_part} - {date}")


def main():
    args = sys.argv[1:]
    root = None
    if "--root" in args:
        i = args.index("--root")
        try:
            root = args[i + 1]
        except IndexError:
            sys.exit("--root requires a directory path")
        del args[i:i + 2]
    if not args:
        sys.exit("Usage: python3 generate_doc.py doc_data.json [output_basename] [--root BUBU_FOLDER]")

    with open(args[0], "r", encoding="utf-8") as fh:
        data = json.load(fh)

    if len(args) >= 2:
        out_basename = os.path.splitext(args[1])[0]
    elif root:
        out_basename = compute_output_basename(data, root)
    else:
        slug = _safe_name(data.get("title") or data.get("doc_type") or "Document")
        date = data.get("date_iso") or datetime.now().strftime("%Y-%m-%d")
        out_basename = f"{slug} - {date}"

    docx_path = build(data, out_basename)
    print(f"DOCX: {docx_path}")


if __name__ == "__main__":
    main()
